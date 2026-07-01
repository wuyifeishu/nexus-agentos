"""
文档加载器 — PDF / DOCX / TXT / Markdown 解析与分块。

零外部 HTTP 依赖，纯本地解析。
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterator


class Document:
    """文档片段。"""

    def __init__(self, content: str, source: str = "", page: int = 0, metadata: dict | None = None):
        self.content = content
        self.source = source
        self.page = page
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document(source={self.source!r}, chars={len(self.content)})"


class DocumentLoader:
    """文档加载器 — 支持多种格式的文档解析与智能分块。

    Args:
        chunk_size: 分块大小（字符数）
        chunk_overlap: 块间重叠字符数
    """

    SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown", ".py", ".json", ".yaml", ".yml"}

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_file(self, path: str) -> list[Document]:
        """加载单个文件，自动识别格式。"""
        path = os.path.abspath(path)
        suffix = Path(path).suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise ValueError(f"不支持的文件格式: {suffix}。支持: {self.SUPPORTED_SUFFIXES}")

        if suffix == ".pdf":
            text = self._read_pdf(path)
        elif suffix == ".docx":
            text = self._read_docx(path)
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        return self._chunk(text, source=path)

    def load_directory(self, dir_path: str, recursive: bool = True) -> list[Document]:
        """加载目录下所有支持的文件。"""
        docs = []
        for root, _, files in os.walk(dir_path):
            for fn in sorted(files):
                fp = os.path.join(root, fn)
                suffix = Path(fp).suffix.lower()
                if suffix in self.SUPPORTED_SUFFIXES:
                    try:
                        docs.extend(self.load_file(fp))
                    except Exception:
                        pass
            if not recursive:
                break
        return docs

    def _read_pdf(self, path: str) -> str:
        """读取 PDF 文本。"""
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError("pypdf 未安装。运行: pip install pypdf")

    def _read_docx(self, path: str) -> str:
        """读取 DOCX 文本。"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx 未安装。运行: pip install python-docx")

    def _chunk(self, text: str, source: str = "") -> list[Document]:
        """固定大小+重叠分块。"""
        if len(text) <= self.chunk_size:
            return [Document(content=text.strip(), source=source)]

        chunks = []
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(Document(content=chunk, source=source))
            start += self.chunk_size - self.chunk_overlap
        return chunks


# ── 便捷函数 ──────────────────────────────────────────────────

def load_file(path: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[Document]:
    """便捷函数：加载单个文件。"""
    loader = DocumentLoader(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return loader.load_file(path)


def load_directory(dir_path: str, recursive: bool = True, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[Document]:
    """便捷函数：加载目录。"""
    loader = DocumentLoader(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return loader.load_directory(dir_path, recursive=recursive)
